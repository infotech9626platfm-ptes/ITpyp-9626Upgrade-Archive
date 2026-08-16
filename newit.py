#************** Development on IT 9626 version starts *********
# ********** Adapted from 9699 Sociology PYP Portal ***********
import datetime
import io
import os
import fitz  # PyMuPDF
import streamlit as st

# Word Document Libraries
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ==========================================
# 0. STREAMLIT PAGE CONFIG & CUSTOM STYLING
# ==========================================
st.set_page_config(
    page_title="9626 IT PYP Portal", 
    page_icon="💻",
    layout="wide"
)

# Custom Styling (Sociology Dashboard Theme)
st.markdown("""
    <style>
    /* 1. Main Page Background */
    .stApp, [data-testid="stAppViewContainer"] {
        background-color: #63D0F8 !important;
    }
    
    /* 2. Sidebar Background */
    [data-testid="stSidebar"], [data-testid="stSidebar"] > div:first-child {
        background-color: #F0FCBB !important;
    }

    /* 3. Global Text Color (#384403) */
    html, body, [class*="css"], h1, h2, h3, h4, h5, h6, p, span, label, div, .stMarkdown {
        color: #384403 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    /* 4. Input Bars, Selectboxes & Text Areas (#A5C809 Border) */
    div[data-baseweb="input"], 
    div[data-baseweb="select"] > div, 
    .stTextInput input, 
    .stSelectbox select,
    textarea {
        background-color: #CEC2F5 !important;
        color: #141801 !important;
        border-radius: 10px !important;
        border: 5px solid #A5C809 !important;
    }

    /* 5. Buttons Styling */
    .stButton button, 
    .stDownloadButton button, 
    [data-testid="baseButton-secondary"], 
    [data-testid="baseButton-primary"] {
        background-color: #C9F40B !important;
        color: #384403 !important;
        border: 2.3px solid #A5C809 !important;
        border-radius: 8px !important;
        font-weight: bold !important;
        transition: all 0.2s ease-in-out;
    }
    
    /* Hover state for buttons */
    .stButton button:hover, .stDownloadButton button:hover {
        background-color: #A5C809 !important;
        color: #384403 !important;
        border: 4px solid #384403 !important;
    }

    /* 6. Navigation Tab Labels */
    button[data-baseweb="tab"] p {
        font-weight: bold !important;
        font-size: 1.2rem !important;
        color: #384403 !important;
    }
    
    /* Active Tab Highlight Indicator */
    div[data-baseweb="tab-highlight"] {
        background-color: #F863E1 !important;
    }

    /* 7. Expanders & Containers */
    [data-testid="stExpander"] {
        border: 1.5px solid #A5C809 !important;
        border-radius: 8px !important;
        background-color: #F0FCBB !important;
    }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 1. DIRECTORY MAPPING & CONFIGURATION
# ==========================================
SYLLABUS_CODE = "9626"
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

# 6 Local Folder Mappings
LOCAL_FOLDERS = {
    "p1_p3": "9626_Paper1n3",
    "p2": "9626_Paper2",
    "p4": "9626_Paper4",
    "zip_files": "9626_ZipfilesP2n4",
    "ms_p1_p3": "9626_ms_P1n3",
    "ms_p2_p4": "9626_ms_P2n4"
}

for folder_path in LOCAL_FOLDERS.values():
    os.makedirs(folder_path, exist_ok=True)

# ==========================================
# 2. SERVICE ACCOUNT AUTHENTICATION & SYNC
# ==========================================
def build_drive_service(write_access=False):
    """Authenticates using Google Service Account credentials from Streamlit Secrets."""
    try:
        if "gcp_service_account" in st.secrets:
            service_account_info = dict(st.secrets["gcp_service_account"])
            scopes = ['https://www.googleapis.com/auth/drive.file'] if write_access else SCOPES
            creds = service_account.Credentials.from_service_account_info(
                service_account_info, 
                scopes=scopes
            )
            return build('drive', 'v3', credentials=creds)
        else:
            st.error("Missing [gcp_service_account] configuration in secrets.")
            return None
    except Exception as e:
        st.error(f"Authentication Error: {e}")
        return None

def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """Downloads missing files from Google Drive folder into local directory."""
    service = build_drive_service(write_access=False)
    if not service:
        return 0, "Failed to authenticate Service Account."
    
    folder_ids = st.secrets.get("drive_folders", {})
    drive_folder_id = folder_ids.get(folder_key)
    
    if not drive_folder_id:
        return 0, f"Missing drive_folder_id for `{folder_key}` in secrets."

    local_path = LOCAL_FOLDERS[folder_key]
    
    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        drive_files = []
        page_token = None

        while True:
            response = service.files().list(
                q=query,
                fields="nextPageToken, files(id, name, mimeType)",
                pageToken=page_token,
                pageSize=100
            ).execute()
            
            drive_files.extend(response.get('files', []))
            page_token = response.get('nextPageToken', None)
            
            if not page_token:
                break

        downloaded_count = 0

        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                downloaded_count += 1

        total_local_files = len([f for f in os.listdir(local_path) if os.path.isfile(os.path.join(local_path, f))])
        return downloaded_count, f"Synced {downloaded_count} new file(s) for `{folder_key}` (Total: {total_local_files})."
        
    except Exception as e:
        return 0, f"Sync error for `{folder_key}`: {e}"

def perform_bulk_sync():
    """Syncs all 6 configured Google Drive folders."""
    total_synced = 0
    messages = []
    for f_key in LOCAL_FOLDERS.keys():
        count, msg = sync_drive_folder_to_local(f_key)
        total_synced += count
        messages.append(msg)
    return total_synced, messages

# ==========================================
# 3. HELPER FUNCTIONS
# ==========================================
def add_page_number_to_run(run):
    """Adds a dynamic Word page number field."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_worksheet_docx(basket_items: list) -> io.BytesIO:
    """Generates a Word document containing selected PDF pages."""
    doc = Document()
    section = doc.sections[0]

    section.page_width = Inches(8.5)
    section.page_height = Inches(11.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("Page ")
    add_page_number_to_run(header_run)

    doc.add_heading(f'PTES {SYLLABUS_CODE} Information Technology Worksheet', level=1)

    for idx, item in enumerate(basket_items):
        doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
        pdf_doc = fitz.open(item['path'])
        page = pdf_doc.load_page(item['page'])
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = io.BytesIO(pix.tobytes("png"))

        doc.add_picture(img_data, width=Inches(7.8), height=Inches(8.8))

        if idx < len(basket_items) - 1:
            doc.add_page_break()
        pdf_doc.close()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_pdf_page_preview(filepath: str, page_num: int = 0):
    """Renders a PDF page to PNG image bytes for preview."""
    try:
        doc = fitz.open(filepath)
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        st.error(f"Unable to render page preview: {e}")
        return None

def execute_pdf_search(folder_key: str, keyword_string: str, paper_filter: str = None) -> list[dict]:
    """Searches PDF files in a specific folder for matching keywords."""
    results = []
    keywords = [k.strip().lower() for k in keyword_string.split(",") if k.strip()]
    folder_path = LOCAL_FOLDERS[folder_key]
    
    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            if file.endswith(".pdf"):
                # Optional paper-specific filter (e.g., matching paper variant in filename)
                if paper_filter and paper_filter not in file.lower():
                    continue
                filepath = os.path.join(folder_path, file)
                try:
                    doc = fitz.open(filepath)
                    for page_num in range(len(doc)):
                        text = doc[page_num].get_text().lower()
                        if all(kw in text for kw in keywords):
                            results.append({
                                "file": file, 
                                "page": page_num, 
                                "path": filepath
                            })
                    doc.close()
                except Exception:
                    continue
    return results

# ==========================================
# 4. SESSION STATE INITIALIZATION
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []

if 'p1_p3_results' not in st.session_state:
    st.session_state.p1_p3_results = []
if 'p2_results' not in st.session_state:
    st.session_state.p2_results = []
if 'p4_results' not in st.session_state:
    st.session_state.p4_results = []

if 'has_auto_synced' not in st.session_state:
    st.session_state.has_auto_synced = True
    with st.spinner("🚀 Waking up portal & auto-syncing IT files via Service Account..."):
        perform_bulk_sync()

# ==========================================
# 5. STREAMLIT UI LAYOUT
# ==========================================
st.title("PUSAT TINGKATAN ENAM SENGKURONG")
st.subheader(f"💻 {SYLLABUS_CODE} Information Technology PYP Resource Library")

# --- SIDEBAR CONTROLS ---
with st.sidebar:
    st.header("🔄 Google Drive Sync")
    if st.button("🔄 Sync Google Drive", use_container_width=True):
        with st.spinner("Syncing Google Drive folders..."):
            synced_count, sync_msgs = perform_bulk_sync()
            st.success(f"Sync Complete! {synced_count} new file(s) downloaded.")
            for m in sync_msgs:
                st.caption(m)

    st.markdown("---")
    st.metric(label="Saved Pages in Cart", value=len(st.session_state.handout_basket))

    if st.button("🗑️ Clear Cart", use_container_width=True):
        st.session_state.handout_basket = []
        st.rerun()

# --- NAVIGATION TABS (7 TABS TOTAL) ---
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "📖 Theory P1/P3", 
    "💻 Practical P2", 
    "⚙️ Practical P4", 
    "🛒 PYP Cart", 
    "📦 ZIP Source", 
    "🔑 Mark_Schm", 
    "🔒 Upload PYP"
])

# --- TAB 1: THEORY PAPERS (PAPER 1 & PAPER 3) ---
with tab1:
    st.subheader("📖 Theory Papers Search (Paper 1 & Paper 3)")
    
    col_p, col_kw = st.columns([1, 2])
    with col_p:
        theory_paper = st.selectbox(
            "Select Theory Component:", 
            ["All Theory Papers", "Paper 1 (AS Theory)", "Paper 3 (A Level Theory)"],
            key="select_theory_paper"
        )
    with col_kw:
        p1_p3_kw = st.text_input(
            "Enter Search Keywords", 
            placeholder="e.g., Normalization, Database, Encryption, Network Topology", 
            key="p1_p3_kw"
        )

    if st.button("Search Theory Papers", key="btn_search_p1_p3"):
        if p1_p3_kw.strip():
            filter_tag = None
            if theory_paper == "Paper 1 (AS Theory)":
                filter_tag = "_qp_1"
            elif theory_paper == "Paper 3 (A Level Theory)":
                filter_tag = "_qp_3"
                
            with st.spinner("Scanning Theory PDFs..."):
                st.session_state.p1_p3_results = execute_pdf_search("p1_p3", p1_p3_kw, filter_tag)
        else:
            st.warning("Please enter at least one keyword.")

    if st.session_state.p1_p3_results:
        st.write(f"Found **{len(st.session_state.p1_p3_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p1_p3_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_p1_p3_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to cart!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_p1_p3_{idx}"
                        )

# --- TAB 2: AS PRACTICAL (PAPER 2) ---
with tab2:
    st.subheader("💻 AS Level Practical Search (Paper 2)")
    
    p2_kw = st.text_input(
        "Enter Keywords", 
        placeholder="e.g., Spreadsheet, VLOOKUP, Database, Query, CSS, HTML", 
        key="p2_kw"
    )

    if st.button("Search Paper 2", key="btn_search_p2"):
        if p2_kw.strip():
            with st.spinner("Scanning Paper 2 PDFs..."):
                st.session_state.p2_results = execute_pdf_search("p2", p2_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.p2_results:
        st.write(f"Found **{len(st.session_state.p2_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p2_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_p2_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to cart!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_p2_{idx}"
                        )

# --- TAB 3: A LEVEL PRACTICAL (PAPER 4) ---
with tab3:
    st.subheader("⚙️ A Level Practical Search (Paper 4)")
    
    p4_kw = st.text_input(
        "Enter Keywords", 
        placeholder="e.g., JavaScript, Animation, Sound Editing, Vector Graphics, 3D Modeling", 
        key="p4_kw"
    )

    if st.button("Search Paper 4", key="btn_search_p4"):
        if p4_kw.strip():
            with st.spinner("Scanning Paper 4 PDFs..."):
                st.session_state.p4_results = execute_pdf_search("p4", p4_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.p4_results:
        st.write(f"Found **{len(st.session_state.p4_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p4_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_p4_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to cart!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_p4_{idx}"
                        )

# --- TAB 4: PYP CART & WORKSHEET GENERATOR ---
with tab4:
    st.subheader("🛒 PYP Cart & Handout Generator")
    
    if len(st.session_state.handout_basket) > 0:
        st.info(f"### 📋 Selected Question Pages ({len(st.session_state.handout_basket)} items)")
        st.caption("Review your selected pages below. Remove any unwanted pages before compiling into Word format.")
        st.markdown("---")
        
        items_to_display = list(st.session_state.handout_basket)
        
        for idx, item in enumerate(items_to_display):
            filename = item.get("file", "Unknown File")
            page_num = item.get("page", 0) + 1
            file_path = item.get("path", "")
            
            with st.expander(f"📄 Item #{idx + 1}: {filename} (Page {page_num})", expanded=False):
                col_preview, col_action = st.columns([3, 1])
                
                with col_preview:
                    if os.path.exists(file_path):
                        img_bytes = render_pdf_page_preview(file_path, item.get("page", 0))
                        if img_bytes:
                            st.image(img_bytes, caption=f"Preview: Page {page_num}", use_container_width=True)
                    else:
                        st.caption(f"Source file path: `{file_path}`")
                
                with col_action:
                    st.markdown("#### Actions")
                    remove_key = f"remove_btn_cart_item_{idx}_{filename}_{page_num}"
                    if st.button("🗑️ Remove Item", key=remove_key, use_container_width=True):
                        st.session_state.handout_basket.pop(idx)
                        st.toast(f"Removed item #{idx + 1} from cart!")
                        st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("### 📝 Export Handout / Worksheet")
        
        with st.spinner("Merging selected pages into Word Worksheet..."):
            doc_buffer = create_worksheet_docx(st.session_state.handout_basket)
            target_filename = f"{SYLLABUS_CODE}_IT_Worksheet.docx"

        st.download_button(
            label="🪄 Download Merged Word Document Worksheet",
            data=doc_buffer,
            file_name=target_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.info("🛒 Your Cart is empty! Search the Theory or Practical tabs and click '➕ Add to Cart' to start building your handout.")

##############################################################################################
# --- TAB 5: PRACTICAL SOURCE FILES (ZIP FILES) ---
with tab5:
    st.subheader("📦 Practical Source Files Repository (Paper 2 & Paper 4 ZIPs)")
    st.caption("Search and download supporting practical files, data files, and assets.")

    zip_folder_path = LOCAL_FOLDERS["zip_files"]
    
    # --- STRUCTURED PARAMETER INPUT BAR ---
    col_level, col_y, col_m, col_v = st.columns([2, 1, 1.5, 1.5])
    
    with col_level:
        target_zip_paper = st.selectbox(
            "Select Qualification Level:", 
            ["All Practical Papers", "Paper 2 (AS)", "Paper 4 (A Level)"],
            key="zip_paper_filter"
        )

    with col_y:
        zip_year = st.text_input(
            "Year (YYYY)", 
            value="2021", 
            placeholder="e.g. 2021", 
            key="zip_year_input"
        )

    with col_m:
        zip_month = st.selectbox(
            "Select Session", 
            [" June (s) ", " November (w) ", " March (m) "], 
            key="zip_mth_select"
        )
        if "June" in zip_month:
            zip_month_code = "s"
        elif "November" in zip_month:
            zip_month_code = "w"
        else:
            zip_month_code = "m"
            
    with col_v:
        # Auto-default variant based on paper selection, or allow manual pick
        default_var_index = 1 if "Paper 4" in target_zip_paper else 0
        zip_variant = st.selectbox(
            "Select Variant", 
            ["02", "04"], 
            index=default_var_index,
            key="zip_var_select"
        )

    # Clean and extract 2-digit year tag (e.g., '2021' -> '21')
    cleaned_zip_year = zip_year.strip()
    short_zip_year = cleaned_zip_year[-2:] if len(cleaned_zip_year) >= 2 else cleaned_zip_year

    # Constructed search identifiers
    search_session_tag = f"{zip_month_code}{short_zip_year}"  # e.g., 's21' or 'm21'
    search_sf_tag = f"sf_{zip_variant}"                      # e.g., 'sf_02' or 'sf_04'
    expected_zip_pattern = f"{SYLLABUS_CODE}_{search_session_tag}_{search_sf_tag}.zip"

    st.markdown("---")

    if os.path.exists(zip_folder_path):
        all_zips = [f for f in os.listdir(zip_folder_path) if f.endswith(('.zip', '.rar', '.7z'))]
        
        filtered_zips = []
        for z_file in all_zips:
            z_lower = z_file.lower()
            
            # Match Level/Paper Filter
            if target_zip_paper == "Paper 2 (AS)" and not ("_02" in z_lower or "_2" in z_lower or "p2" in z_lower):
                continue
            if target_zip_paper == "Paper 4 (A Level)" and not ("_04" in z_lower or "_4" in z_lower or "p4" in z_lower):
                continue

            # Match Session (e.g., 's21') and Variant Tag (e.g., 'sf_02')
            if search_session_tag in z_lower and search_sf_tag in z_lower:
                filtered_zips.append(z_file)
            # Fallback: match if year tag and variant are both present
            elif search_session_tag in z_lower and zip_variant in z_lower:
                if z_file not in filtered_zips:
                    filtered_zips.append(z_file)

        if filtered_zips:
            st.success(f"Found **{len(filtered_zips)}** matching Practical Source File(s):")
            for z_file in sorted(filtered_zips):
                c_name, c_btn = st.columns([3, 1])
                full_z_path = os.path.join(zip_folder_path, z_file)
                size_mb = os.path.getsize(full_z_path) / (1024 * 1024)
                
                with c_name:
                    st.write(f"📦 **{z_file}** ({size_mb:.2f} MB)")
                with c_btn:
                    with open(full_z_path, "rb") as f_zip:
                        st.download_button(
                            label=f"📥 Download ZIP",
                            data=f_zip,
                            file_name=z_file,
                            mime="application/zip",
                            key=f"dl_zip_{z_file}"
                        )
        else:
            st.warning(f"No ZIP source file found matching session `{search_session_tag}` and variant `{zip_variant}` (Expected pattern: `{expected_zip_pattern}`).")
            st.info("💡 **Tip**: Click **Sync Google Drive** from the sidebar to fetch newly uploaded source files.")
    else:
        st.warning("ZIP directory does not exist.")
###############################################################################################
# --- TAB 6: ANSWER SCHEMES ---
with tab6:
    st.subheader("🔑 Download Marking Schemes")
    
    col_paper, col_y, col_m, col_v = st.columns([2, 1, 1.5, 1.5])
    
    with col_paper:
        target_paper_selection = st.selectbox(
            "Select Target Paper Component:", 
            ["Paper 1", "Paper 2", "Paper 3", "Paper 4"],
            key="ms_paper_select"
        )
        
        # Route to appropriate local folder based on selected paper
        if target_paper_selection in ["Paper 1", "Paper 3"]:
            ms_folder_key = "ms_p1_p3"
        else:
            ms_folder_key = "ms_p2_p4"

    with col_y:
        as_year = st.text_input(
            "Year", 
            value="2021", 
            placeholder="e.g. 2021", 
            key="ms_year"
        )

    with col_m:
        as_month = st.selectbox(
            "Session", 
            [" June (s) ", " November (w) ", " March (m) "], 
            key="ms_mth"
        )
        if "June" in as_month:
            month_code = "s"
        elif "November" in as_month:
            month_code = "w"
        else:
            month_code = "m"
            
    with col_v:
        paper_num = target_paper_selection.split(" ")[1]  # Extract 1, 2, 3, or 4
        default_variants = [f"{paper_num}1", f"{paper_num}2", f"{paper_num}3"]
        as_variant = st.selectbox(
            "Select Variant", 
            default_variants, 
            key="ms_var"
        )

    cleaned_year = as_year.strip()
    short_year = cleaned_year[-2:] if len(cleaned_year) >= 2 else cleaned_year

    search_session_tag = f"{month_code}{short_year}"
    expected_ms_filename = f"{SYLLABUS_CODE}_{search_session_tag}_ms_{as_variant}.pdf"

    st.markdown("---")
    
    found_ms_files = []
    folder_path = LOCAL_FOLDERS[ms_folder_key]

    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            if file.endswith(".pdf"):
                file_lower = file.lower()
                if search_session_tag in file_lower and f"ms_{as_variant}" in file_lower:
                    found_ms_files.append(os.path.join(folder_path, file))
                elif search_session_tag in file_lower and "ms" in file_lower and as_variant in file_lower:
                    if os.path.join(folder_path, file) not in found_ms_files:
                        found_ms_files.append(os.path.join(folder_path, file))

    if found_ms_files:
        st.success(f"Found {len(found_ms_files)} matching Answer Scheme file(s):")
        
        for ms_path in found_ms_files:
            ms_filename = os.path.basename(ms_path)
            
            with st.expander(f"🔑 Mark Scheme: {ms_filename}", expanded=True):
                col_dl, _ = st.columns([1, 2])
                with col_dl:
                    with open(ms_path, "rb") as f:
                        st.download_button(
                            label=f"📥 Download {ms_filename}",
                            data=f,
                            file_name=ms_filename,
                            mime="application/pdf",
                            key=f"dl_ms_btn_{ms_filename}"
                        )
                
                doc = fitz.open(ms_path)
                total_pages = len(doc)
                st.caption(f"📜 Showing all **{total_pages}** pages below:")
                
                # --- SCROLLABLE CONTAINER ---
                with st.container(height=650):
                    for page_num in range(total_pages):
                        page_img = render_pdf_page_preview(ms_path, page_num)
                        if page_img:
                            st.image(
                                page_img, 
                                caption=f"Page {page_num + 1} of {total_pages}", 
                                use_container_width=True
                            )
                            if page_num < total_pages - 1:
                                st.markdown("---")
                        
                doc.close()
    else:
        st.warning(f"No Mark Scheme found matching session `{search_session_tag}` and variant `{as_variant}` (Expected pattern: `{expected_ms_filename}`).")
        st.info("💡 **Tip**: Click **Sync Google Drive** from the sidebar to fetch updated mark scheme files into local storage.")

# --- TAB 7: ADMIN PANEL (6 GOOGLE DRIVE FOLDERS) ---
with tab7:
    st.subheader("🔒 Administrator Control Panel")
    st.caption("Manage Google Drive repositories across all 6 subject partitions.")

    admin_pwd = st.secrets.get("ADMIN_PASSWORD", "")
    pwd_input = st.text_input("Enter Admin Password", type="password", key="admin_pwd_input")

    if pwd_input and pwd_input == admin_pwd:
        st.success("Authenticated as Administrator")
        st.markdown("---")
        
        st.markdown("### 🌐 Google Drive Web Repositories")
        st.info("💡 Click any button below to open its respective Google Drive folder in a new tab where you can upload new files.")
        
        drive_links = st.secrets.get("drive_web_links", {})

        c1, c2, c3 = st.columns(3)
        with c1:
            st.link_button("📖 1. Papers 1 & 3 (Theory)", drive_links.get("p1_p3", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("📦 4. Practical Source Files (ZIP)", drive_links.get("zip_files", "https://drive.google.com"), use_container_width=True)
            
        with c2:
            st.link_button("💻 2. Paper 2 (AS Practical)", drive_links.get("p2", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🔑 5. Answer Schemes (P1 & P3)", drive_links.get("ms_p1_p3", "https://drive.google.com"), use_container_width=True)
            
        with c3:
            st.link_button("⚙️ 3. Paper 4 (A Practical)", drive_links.get("p4", "https://drive.google.com"), use_container_width=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🔑 6. Answer Schemes (P2 & P4)", drive_links.get("ms_p2_p4", "https://drive.google.com"), use_container_width=True)

    elif pwd_input:
        st.error("Incorrect Admin Password.")

# ==========================================
# 6. PORTAL FOOTER
# ==========================================
st.markdown("---")
SCHOOL_NAME = "Pusat Tingkatan Enam Sengkurong (PTES)"
SCHOOL_VISION = "Nurturing Resilient Leaders & Future-Ready Citizens"

footer_html = f"""
<div style="text-align: center; padding: 15px 0px; font-family: sans-serif;">
    <p style="margin: 0; font-size: 1.0em; font-weight: bold; color: #384403;">🏫 {SCHOOL_NAME}</p>
    <p style="margin: 5px 0; font-size: 0.9em; font-style: italic; color: #384403;">"{SCHOOL_VISION}"</p>
    <p style="margin: 5px 0 0 0; font-size: 0.85em; font-weight: 600; color: #384403;">💻 Developed for Computer Science & IT Department ({SYLLABUS_CODE})</p>
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)
